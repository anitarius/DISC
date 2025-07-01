import streamlit as st
import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns
import io
from PIL import Image 
from docx.shared import Pt, RGBColor, Inches


def frontend():
    image_path = "portada.jpeg"  # Debe estar en la misma carpeta que app.py
    try:
        portada = Image.open(image_path)
        st.image(portada, use_column_width=True)
    except FileNotFoundError:
        st.warning(" No se encontró la imagen 'portada.jpg'. Asegurate de que esté en la misma carpeta que app.py.")

    st.set_page_config(page_title="DISC AUTO", layout="centered")
    st.title("DISC AUTO")
    st.markdown(
        """ 
        1) Subir el archivo a evaluar: excel con el formato prestablecido.
        2) Descargar archivo procesado.
        """
    )

    uploaded_file = st.file_uploader("Subir archivo de Excel", type=["xlsx"])
    if uploaded_file is not None:
        entrada = uploaded_file.name
        st.write(f"Archivo subido: {entrada}")
        procesar_archivo(uploaded_file)

def procesar_archivo(uploaded_file):
    
    # Carga y prerocesamiento del archivo
    df = pd.read_excel(uploaded_file)
    df = df.replace('X', 1)
    df = df.fillna(0)
    nombre_row_index = df[df.iloc[:, 0] == 'Nombre:'].index[0]
    edad_row_index = df[df.iloc[:, 0] == 'Edad:'].index[0]
    fecha_row_index = df[df.iloc[:, 0] == 'Fecha  :'].index[0]


    nombre = df.iloc[nombre_row_index, 1]
    edad = df.iloc[edad_row_index, 1]
    fecha = df.iloc[fecha_row_index, 1]
    #print(f"\nNombre: {nombre}, Edad: {edad}, Fecha: {fecha}")

    rows_to_drop = [0, 1, 2, 3, 4] #10, 15, 20, 25, 30, 35]
    df = df.drop(rows_to_drop)

    #print("Archivo de Excel leído exitosamente.")


    #1) sumar
    col_map = {
        'B': 1, 'C': 2,  # primera palabra
        'E': 4, 'F': 5,  # segunda palabra
        'H': 7, 'I': 8,  # tercera palabra
        'K': 10, 'L': 11 # cuarta palabra
    }

    coordenadas = {
    'D+': [('B',2),('B',6),('B',11),('B',16),('B',17),('B',24),('B',27),
           ('E',4),('E',8),('E',9),('E',15),('E',20),('E',21),('E',27),
           ('H',3),('H',6),('H',11),('H',14),('H',17),('H',24),('H',28),
           ('K',3),('K',7),('K',11),('K',13),('K',19),('K',21),('K',26)],

    'D-': [('C',2),('C',6),('C',11),('C',16),('C',17),('C',24),('C',27),
           ('F',4),('F',8),('F',9),('F',15),('F',20),('F',21),('F',27),
           ('I',3),('I',6),('I',11),('I',14),('I',17),('I',24),('I',28),
           ('L',3),('L',7),('L',11),('L',13),('L',19),('L',21),('L',26)],

    'I+': [('B',1),('B',7),('B',9),('B',13),('B',19),('B',22),('B',25),
           ('E',1),('E',7),('E',10),('E',16),('E',17),('E',23),('E',28),
           ('H',1),('H',8),('H',9),('H',16),('H',18),('H',21),('H',27),
           ('K',1),('K',5),('K',9),('K',16),('K',17),('K',23),('K',27)],

    'I-': [('C',1),('C',7),('C',9),('C',13),('C',19),('C',22),('C',25),
           ('F',1),('F',7),('F',10),('F',16),('F',17),('F',23),('F',28),
           ('I',1),('I',8),('I',9),('I',16),('I',18),('I',21),('I',27),
           ('L',1),('L',5),('L',9),('L',16),('L',17),('L',23),('L',27)],

    'S+': [('B',4),('B',8),('B',12),('B',15),('B',20),('B',21),('B',28),
           ('E',3),('E',6),('E',11),('E',14),('E',18),('E',22),('E',26),
           ('H',4),('H',7),('H',10),('H',13),('H',19),('H',22),('H',26),
           ('K',4),('K',8),('K',10),('K',15),('K',20),('K',22),('K',28)],

    'S-': [('C',4),('C',8),('C',12),('C',15),('C',20),('C',21),('C',28),
           ('F',3),('F',6),('F',11),('F',14),('F',18),('F',22),('F',26),
           ('I',4),('I',7),('I',10),('I',13),('I',19),('I',22),('I',26),
           ('L',4),('L',8),('L',10),('L',15),('L',20),('L',22),('L',28)],

    'C+': [('B',3),('B',5),('B',10),('B',14),('B',18),('B',23),('B',26),
           ('E',2),('E',5),('E',12),('E',13), ('E',19),('E',24),('E',25),
           ('H',2),('H',5),('H',12),('H',15),('H',20),('H',23),('H',25),
           ('K',2),('K',6),('K',12),('K',14),('K',18),('K',24),('K',25)],

    'C-': [('C',3),('C',5),('C',10),('C',14),('C',18),('C',23),('C',26),
           ('F',2),('F',5),('F',12),('F',13),('F',19),('F',24),('F',25),
           ('I',2),('I',5),('I',12),('I',15),('I',20),('I',23),('I',25),
           ('L',2),('L',6),('L',12),('L',14),('L',18),('L',24),('L',25)]
}

    indexados = {k: [(fila-1, col_map[col]) for col, fila in v] for k, v in coordenadas.items()}

    def sumar_posiciones(df, index_dict):
        resultados = {}
        for key, coords in index_dict.items():
            suma = sum(df.iat[fila, col] for fila, col in coords if 0 <= fila < df.shape[0])
            resultados[key] = suma
        return resultados

    res = sumar_posiciones(df, indexados)
    resumen = pd.DataFrame({
        'DISC': ['D', 'I', 'S', 'C'],
        'Positivo': [res['D+'], res['I+'], res['S+'], res['C+']],
        'Negativo': [res['D-'], res['I-'], res['S-'], res['C-']]
    })
    resumen['Neto'] = resumen['Positivo'] - resumen['Negativo']

    #print(resumen)

    #2) crear categorizacion
    x1 = [0, 0, 0, 0]
    conditions = [
        [(0, 2, 1), (3, 4, 2), (5, 6, 3), (7, 7, 4), (8, 9, 5), (10, 12, 6), (13, float('inf'), 7)],
        [(0, 2, 1), (3, 3, 2), (4, 5, 3), (6, 6, 4), (7, 7, 5), (8, 9, 6), (10, float('inf'), 7)],
        [(0, 2, 1), (3, 3, 2), (4, 4, 3), (5, 5, 4), (6, 6, 5), (7, 8, 6), (9, float('inf'), 7)],
        [(0, 3, 1), (4, 4, 2), (5, 5, 3), (6, 7, 4), (8, 8, 5), (9, 10, 6), (11, float('inf'), 7)]]
    for i in range(len(conditions)):
        for lower, upper, value in conditions[i]:
            if lower <= resumen['Positivo'][i] <= upper:
                x1[i] = value
                break

    x2 = [0, 0, 0, 0]

    conditions = [
        [(0, 1, 7), (2, 3, 6), (4, 5, 5), (6, 6, 4), (7, 8, 3), (9, 10, 2), (11, float('inf'), 1)],
        [(1, 3, 7), (4, 5, 6), (6, 6, 5), (7, 7, 4), (8, 8, 3), (9, 10, 2), (11, float('inf'), 1)],
        [(0,3,7) , (4,4,6), (5,6,5),(7,7,4),(8,9,3),(10,11,2),(12,float('inf'),1)],
        [(0,2,7),(3,3,6),(4,4,5),(5,5,4),(6,7,3),(8,8,2),(9,float('inf'),1)]]

    for i in range(len(conditions)):
        for lower, upper, value in conditions[i]:
            if lower <= resumen['Negativo'][i] <= upper:
                x2[i] = value
                break

    x3 = [0, 0, 0, 0]

    conditions = [
        [(10, 28, 7), (6, 9, 6), (2, 5, 5), (0, 1, 4), (-3, -1, 3), (-6, -4, 2), (-float('inf'), -7, 1)],
        [(5, 28, 7), (3, 4, 6), (0, 2, 5), (-2, -1, 4), (-4, -3, 3), (-7, -5, 2), (-float('inf'), -8, 1)],
        [(5, 28, 7), (2, 4, 6), (-1, 1, 5), (-3, -2, 4), (-6, -4, 3), (-9, -7, 2), (-float('inf'), -10, 1)],
        [(8, 28, 7), (5, 7, 6), (2, 4, 5), (0, 1, 4), (-2, -1, 3), (-5, -3, 2), (-float('inf'), -6, 1)]
    ]
    for i in range(len(conditions)):
        for lower, upper, value in conditions[i]:
            if lower <= resumen['Neto'][i] <= upper:
                x3[i] = value
                break

    #print('x2:',x2)
    #print('x1:', x1)
    #print('x3:', x3)

    #4) interpretar
    excel_file_interpretacion_path = 'Test DISC (tablas interpretación).xlsx'
    df_interpretacion = pd.DataFrame()
    with open(excel_file_interpretacion_path, 'rb') as f:
        df_interpretacion = pd.read_excel(excel_file_interpretacion_path)
    #print(f"Error: El archivo de tablas de interpretación no se encontró en la ruta especificada: {excel_file_interpretacion_path}")

    clave_x1_str = "".join(map(str, x1))
    clave_x1_int = int(clave_x1_str)

    clave_x2_str = "".join(map(str, x2))
    clave_x2_int = int(clave_x2_str)

    clave_x3_str = "".join(map(str, x3))
    clave_x3_int = int(clave_x3_str)

    interpretacion_x1 = df_interpretacion[df_interpretacion['clave'] == clave_x1_int]['df'].values
    interpretacion_x2 = df_interpretacion[df_interpretacion['clave'] == clave_x2_int]['df'].values
    interpretacion_x3 = df_interpretacion[df_interpretacion['clave'] == clave_x3_int]['df'].values


    if len(interpretacion_x1) > 0:
        valor = interpretacion_x1[0]
    # print(f"The corresponding value from df_interpretacion for {clave_x1_int} is: {valor}")
    else:
        #print(f"No matching key found in df_interpretacion for {clave_x1_int}")
        valor = None
    if len(interpretacion_x2) > 0:
        valor = interpretacion_x2[0]
    # print(f"The corresponding value from df_interpretacion for {clave_x2_int} is: {valor}")
    else:
        #print(f"No matching key found in df_interpretacion for {clave_x2_int}")
        valor = None
    if len(interpretacion_x3) > 0:
        valor = interpretacion_x3[0]
        #print(f"The corresponding value from df_interpretacion for {clave_x3_int} is: {valor}")
    else:
        #print(f"No matching key found in df_interpretacion for {clave_x3_int}")
        valor = None

    #ARMAR documento
    document = Document()
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Change document margin
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # INSERTAR LOGO EN EL ENCABEZADO (alineado a la derecha)
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_picture('logo.png', width=Inches(1.9))

# Espaciado después del encabezado
    document.add_paragraph()
# TÍTULO PRINCIPAL (¡Hola Roberta!)
    titulo = document.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = titulo.add_run(f"¡Hola {nombre}!")
    run.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    run.font.size = Pt(24)
    run.font.name = "Arial"
    
# Texto introductorio
    p1 = document.add_paragraph()
    p1.add_run(
    "A continuación, verás el resultado de tu test "
    )
    p1.add_run("DISC.\n").bold = True
    p1.add_run("En síntesis, esta prueba mide cómo hacemos las cosas y cómo nos relacionamos con los demás.\n"
    "Nos brinda información sobre cómo es nuestro estilo en tres situaciones: el estilo que tenemos "
    "de comportamiento diario o integral (el que ponemos en juego cuando nos desenvolvemos cotidianamente "
    "en el mundo), el estilo natural o de motivación y el estilo adaptado ante situaciones de tensión."
    )
    for run in p1.runs:
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(44, 62, 80)  # Azul oscuro
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    import seaborn as sns
    
    #primer grafico
    disc_labels = ['D','I','S','C']

    #titulo
    titulo = document.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f"PERFIL INTEGRAL (COMPORTAMIENTO DIARIO): {interpretacion_x3[0]}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Aptos"
    plt.figure(figsize=(6, 5))
    colors = ['#FF9999', '#FFFF99', '#99FF99', '#99CCFF'] 
    sns.barplot(x=disc_labels, y=x3, palette=colors)

    plt.ylim(0, 8)
    for i, score in enumerate(x3):
        plt.hlines(score, i - 0.4, i + 0.4, color='black', linestyles='dashed')
    plt.savefig('x3_plot.png')
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    document.add_picture('x3_plot.png', width=Inches(5)) 
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Espaciado después del encabezado
    document.add_page_break()

    document.add_paragraph()

    titulo = document.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f"PERFIL DE MOTIVACIÓN: {interpretacion_x1[0]}           PERFIL DE BAJO PRESIÓN: {interpretacion_x2[0]}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Aptos"

    fig, axes = plt.subplots(1, 2, figsize=(10, 7)) 

    colors = ['#FF9999', '#FFFF99', '#99FF99', '#99CCFF'] 

    # Plot X1 on the left subplot
   
    sns.barplot(x=disc_labels, y=x1, palette=colors, ax=axes[0])
    axes[0].set_title('Perfil de motivación')
    axes[0].set_ylim(0, 8)
    for i, score in enumerate(x1):
        axes[0].hlines(score, i - 0.4, i + 0.4, color='black', linestyles='dashed')

    # Plot X2 on the right subplot
    sns.barplot(x=disc_labels, y=x2, palette=colors, ax=axes[1])
    axes[1].set_title('Perfil bajo presión')
    axes[1].set_ylim(0, 8)
    for i, score in enumerate(x2):
        axes[1].hlines(score, i - 0.4, i + 0.4, color='black', linestyles='dashed')

    plt.tight_layout() 
    plt.savefig('x1_x2_plots.png')

    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    document.add_picture('x1_x2_plots.png', width=Inches(6.7)) # Adjust width as needed to fit increased figure size
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #Perfil de comportamiento diario
    from docx import Document
    import os
    document_to_append_path = f'/Diario/{interpretacion_x3[0]}.docx'

    if os.path.exists(document_to_append_path):
        doc_to_append = Document(document_to_append_path)

        # Agrega cada elemento del documento a insertar
        for element in doc_to_append.element.body:
            document.element.body.append(element)
    else:
        print(f"❌ Error: No se encontró el documento en la ruta: {document_to_append_path}")
    
    #salida
    salida = st.text_input("✏️ Elegí un nombre para el archivo Word", value="informe_DISC")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="⬇️ Descargar informe Word",
        data=buffer,
        file_name=f"{salida}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

frontend()
